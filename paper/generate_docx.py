"""
Generate acm_paper.docx and elsevier_paper.docx from paper content.
Run: python generate_docx.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_two_columns(doc):
    """Apply two-column layout to the entire document body."""
    section = doc.sections[0]
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')  # 0.5 inch gap
    cols.set(qn('w:equalWidth'), '1')
    sectPr.append(cols)


def set_margins(doc, top=1.0, bottom=1.0, left=0.75, right=0.75):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)


def add_heading(doc, text, level=1, font_name='Times New Roman',
                font_size=12, bold=True, space_before=10, space_after=4,
                color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    return p


def add_body(doc, text, font_name='Times New Roman', font_size=10,
             space_before=0, space_after=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.italic = italic
    return p


def add_mixed(doc, parts, font_name='Times New Roman', font_size=10,
              space_before=0, space_after=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """parts: list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = font_name
        run.font.size = Pt(font_size)
    return p


def shade_table_header(row, color='D9D9D9'):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)


def style_table(table, font_name='Times New Roman', font_size=9):
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                for run in p.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)


def add_rule(doc):
    """Add a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ---------------------------------------------------------------------------
# ACM Paper
# ---------------------------------------------------------------------------

def build_acm():
    doc = Document()
    set_margins(doc, top=1.0, bottom=1.0, left=0.75, right=0.75)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    # ---- Title ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(
        'Weaponized Weights: A Cryptographic Provenance and Binary Inspection '
        'Framework for Securing Machine Learning Model Artifacts in Enterprise '
        'MLOps Pipelines'
    )
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)

    # ---- Authors ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('Sunil Gentyala')
    r.bold = True; r.font.size = Pt(11); r.font.name = 'Times New Roman'

    for line in [
        'HCLTech, Cybersecurity and AI Security Practice',
        'Dallas, TX, USA',
        'sunil.gentyala@ieee.org',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(line)
        r.font.size = Pt(10); r.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('Rakesh Prakash')
    r.bold = True; r.font.size = Pt(11); r.font.name = 'Times New Roman'

    for line in [
        'University of Colorado Boulder, Department of Computer Science',
        'Boulder, CO, USA',
        'Rakesh.Prakash@colorado.edu',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(line)
        r.font.size = Pt(10); r.font.name = 'Times New Roman'

    doc.add_paragraph()

    # ---- CCS + Keywords box ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run('CCS Concepts: ')
    r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    r2 = p.add_run(
        'Security and privacy → Software security engineering; '
        'Security and privacy → Malware and its mitigation; '
        'Computing methodologies → Machine learning.'
    )
    r2.font.size = Pt(9); r2.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run('Keywords: ')
    r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    r2 = p.add_run(
        'machine learning security, supply chain attack, model provenance, '
        'Safetensors, GGUF, pickle deserialization, MLOps, cryptographic '
        'verification, CI/CD security, binary inspection'
    )
    r2.font.size = Pt(9); r2.font.name = 'Times New Roman'

    add_rule(doc)

    # Enable two-column layout after the title block
    set_two_columns(doc)

    # ---- Abstract ----
    add_heading(doc, 'ABSTRACT', level=1, font_size=10, space_before=8)
    add_body(doc,
        'The proliferation of publicly distributed machine learning model artifacts '
        'has introduced a structural security risk that conventional enterprise '
        'tooling is fundamentally ill-equipped to address. Weight files serialized in '
        'PyTorch checkpoint, Safetensors, and GGUF formats are routinely downloaded '
        'from public model repositories and admitted into production inference '
        'pipelines without binary-level inspection or cryptographic provenance '
        'verification. Existing static analysis, dynamic analysis, and antivirus '
        'infrastructure have no visibility into these binary weight formats, leaving '
        'organizations exposed to a class of supply chain attack that operates '
        'entirely below the application layer.',
        space_before=4)
    add_body(doc,
        'We present model-provenance-guard, a production-ready, open-source toolkit '
        'comprising a Python binary inspector (verify_weights.py) and a five-stage '
        'GitHub Actions provenance gate. The pipeline enforces SHA-256 hash '
        'verification against an internal trusted registry, Sigstore/Cosign '
        'cryptographic signature checking, pickle opcode scanning via picklescan, '
        'Safetensors header anomaly detection, and GGUF metadata key auditing before '
        'any artifact is promoted within a downstream pipeline. We evaluate the toolkit '
        'against eight synthetic adversarial test cases spanning dtype manipulation, '
        'data offset injection, magic byte substitution, and malformed JSON header '
        'attacks. The inspector achieved zero false negatives and zero false positives '
        'across all adversarial scenarios, demonstrating that shift-left provenance '
        'enforcement is both technically feasible and operationally practical at the '
        'CI/CD boundary.',
        space_before=4)

    # ---- 1. Introduction ----
    add_heading(doc, '1  INTRODUCTION', font_size=11, space_before=12)
    add_body(doc,
        'The security community\'s attention toward artificial intelligence systems '
        'has concentrated heavily on prompt injection and model output manipulation '
        'over the past two years. Those concerns are legitimate, but their prominence '
        'has allowed a structurally more dangerous threat to mature with insufficient '
        'scrutiny. Adversaries are no longer confined to crafting adversarial inputs '
        'at runtime. They are embedding threats directly inside compiled weight '
        'artifacts that machine learning engineers download from public repositories '
        'as a routine operational act.')
    add_body(doc,
        'The attack surface has migrated from the application layer to the binary '
        'layer, and the tooling most enterprises rely on for security assurance lacks '
        'any meaningful visibility there. Static application security testing (SAST) '
        'tools analyze source code. Dynamic application security testing (DAST) tools '
        'probe running applications. Antivirus engines match file content against '
        'known malware signatures. None of these approaches has any concept of what '
        'lies inside a 4-bit quantized GGUF binary, a Safetensors archive, or a '
        'PyTorch checkpoint serialized with Python\'s pickle protocol.')
    add_body(doc,
        'Three weight file formats account for the substantial majority of model '
        'artifacts in contemporary enterprise inference deployments. PyTorch .pt and '
        '.pth checkpoints use Python\'s pickle serialization protocol, which by '
        'design permits arbitrary Python objects to be deserialized, creating a '
        'direct remote code execution (RCE) vector at the moment torch.load() is '
        'called [2]. Safetensors was introduced specifically to eliminate '
        'deserialization-time code execution; however, its JSON header presents '
        'its own attack surface [8]. GGUF, the dominant format for quantized '
        'inference, includes a flexible metadata section with no schema enforcement, '
        'and multiple heap-buffer-overflow vulnerabilities in its reference parser '
        'were disclosed in early 2024 [11].')
    add_body(doc,
        'This paper makes the following contributions: (1) A format-level threat '
        'model covering PyTorch, Safetensors, and GGUF attack surfaces. '
        '(2) verify_weights.py, a binary inspector that performs header anomaly '
        'detection and metadata key auditing without loading tensor data. '
        '(3) A five-stage GitHub Actions provenance gate integrating registry hash '
        'verification, Sigstore/Cosign signature checking, picklescan analysis, and '
        'binary inspection. (4) An empirical evaluation against eight synthetic '
        'adversarial test cases with zero false negatives and zero false positives. '
        '(5) An open-source release at '
        'https://github.com/sunilgentyala/model-provenance-guard.')

    # ---- 2. Background ----
    add_heading(doc, '2  BACKGROUND AND RELATED WORK', font_size=11, space_before=12)

    add_heading(doc, '2.1  Machine Learning Supply Chain Attacks', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The concept of a software supply chain attack is well established in the '
        'security literature, but its application to machine learning weight artifacts '
        'is comparatively recent. Jiang et al. [4] conducted a systematic measurement '
        'study of malicious code poisoning attacks on pre-trained model hubs and found '
        'that attackers exploit the implicit trust engineers place in popular repository '
        'namespaces. Documented incidents on Hugging Face have confirmed that malicious '
        'model artifacts can execute reverse shell payloads at load time. The '
        'Communications of the ACM survey [5] argues that the combination of high '
        'download volumes and weak provenance controls makes model hubs uniquely '
        'attractive targets relative to traditional software package registries.')
    add_body(doc,
        'The zero-trust perspective on AI data pipelines, articulated by Mudusu and '
        'Gentyala [1], establishes that the never-trust-always-verify principle must '
        'be applied directly to data and model artifacts throughout the pipeline. '
        'That work demonstrates that cryptographic authentication and tamper-evident '
        'lineage tracking can achieve complete detection of data tampering and '
        'component impersonation in production AI systems. Our framework extends '
        'those principles specifically to weight file artifacts at the CI/CD intake '
        'boundary.')

    add_heading(doc, '2.2  Pickle Deserialization in PyTorch Checkpoints', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The vulnerability class most thoroughly studied in the context of ML weight '
        'files is pickle deserialization. Python\'s pickle module supports a '
        '__reduce__ protocol that enables arbitrary Python callables to execute '
        'during deserialization. Any torch.load() call on a maliciously constructed '
        'checkpoint can trigger operating system command execution without any user '
        'interaction [2]. Koishybayev et al. [2] present PickleBall, a secure '
        'deserialization system that intercepts pickle opcodes before object '
        'reconstruction and enforces a safe-by-default allowlist.')

    add_heading(doc, '2.3  Safetensors Format and Header Security', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The Safetensors format was developed by Hugging Face as a direct response '
        'to the RCE risks inherent in pickle-based serialization [8]. By replacing '
        'pickle with a JSON-encoded header followed by raw tensor byte data, '
        'Safetensors eliminates arbitrary code execution at load time. However, the '
        'header parsing step introduces its own attack surface through oversized '
        'JSON headers, malformed Unicode sequences, invalid tensor dtype values, and '
        'data offsets that exceed file boundaries. CryptoTensors [9] proposes '
        'extending this format with tensor-level encryption and embedded access '
        'control policies.')

    add_heading(doc, '2.4  GGUF Format Vulnerabilities', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'GGUF, introduced in August 2023 and adopted as the standard distribution '
        'format for quantized models in the Ollama and llama.cpp ecosystems, '
        'presents a distinct set of security concerns. Cisco Talos disclosed three '
        'heap-buffer-overflow vulnerabilities in llama.cpp\'s GGUF parser in early '
        '2024, arising from insufficient bounds checking in metadata key-value '
        'parsing. Additionally, researchers have demonstrated that quantization error '
        'provides sufficient numerical flexibility to construct adversarial models '
        'whose malicious behavior is invisible in full-precision analysis [11].')

    add_heading(doc, '2.5  Neural Backdoor Attacks', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The neural backdoor threat, formalized by Gu et al. in the seminal BadNets '
        'paper [3], involves poisoning model weights during training such that the '
        'model produces correct outputs on clean inputs but triggers a targeted '
        'misclassification on inputs containing a specific pattern. This attack class '
        'is undetectable by binary-level inspection: the weight file loads cleanly, '
        'passes format validation, and produces correct behavior under standard test '
        'inputs.')

    add_heading(doc, '2.6  Software Artifact Signing', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'Sigstore [6] provides a keyless signing infrastructure that ties artifact '
        'signatures to ephemeral OIDC identity tokens, dramatically lowering the '
        'adoption barrier relative to traditional GPG-based signing. As of 2025, '
        'Sigstore underpins artifact signing for npm, PyPI, and Kubernetes, and the '
        'OpenSSF AI/ML Working Group has adapted these mechanisms specifically for '
        'model signing [15].')

    add_heading(doc, '2.7  ML Provenance and Governance', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'Schlegel and Sattler [7] present a PROV-compliant provenance model for '
        'end-to-end ML pipelines that enables tamper-evident audit trails. '
        'Gentyala [12] frames the AI supply chain problem in terms of AI Software '
        'Bills of Materials, demonstrating that CycloneDX and SPDX can encode model '
        'weight metadata and cryptographic attestations. The NIST AI Risk Management '
        'Framework [13] and MITRE ATLAS [14] provide complementary governance '
        'scaffolding. The OWASP Top 10 for LLM Applications explicitly names supply '
        'chain compromise as a critical risk category [16].')

    # ---- 3. Threat Model ----
    add_heading(doc, '3  THREAT MODEL', font_size=11, space_before=12)

    add_heading(doc, '3.1  Adversary Model', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'We consider a threat actor who has obtained write access to a model '
        'repository on a public model hub through account compromise of a legitimate '
        'model publisher, submission of a malicious pull request, or creation of a '
        'confusingly named fork. The adversary\'s goal is to deliver a weaponized '
        'weight artifact to at least one GPU-enabled inference endpoint inside a '
        'target enterprise environment. We assume the target organization\'s CI/CD '
        'pipeline downloads model artifacts automatically from external sources as '
        'part of its standard model update workflow.')

    add_heading(doc, '3.2  PyTorch Checkpoint Attack Surface', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'PyTorch checkpoints use Python\'s pickle serialization protocol. The pickle '
        'module supports a __reduce__ method that, when present on a serialized '
        'object, is called by the deserializer to reconstruct the object. An attacker '
        'can craft a checkpoint that includes an object whose __reduce__ method '
        'returns a callable paired with arguments, effectively embedding an arbitrary '
        'function call that executes the moment torch.load() is invoked. Common '
        'payloads include reverse shells, credential harvesters, and persistence '
        'mechanisms.')

    add_heading(doc, '3.3  Safetensors Attack Surface', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'While the Safetensors format eliminates deserialization-time code execution, '
        'it introduces three distinct attack surfaces. First, the JSON header is '
        'vulnerable to parser exploits through oversized headers and deeply nested '
        'structures. Second, the __metadata__ dictionary accepts arbitrary string '
        'key-value pairs with no schema enforcement. Third, and most consequentially, '
        'Safetensors files can carry weights poisoned at training time [3]. The file '
        'loads cleanly under all format-level checks; the backdoor activates only on '
        'attacker-chosen inputs.')

    add_heading(doc, '3.4  GGUF Attack Surface', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'GGUF\'s binary metadata section has no schema enforcement beyond type '
        'tagging. A threat actor can embed arbitrary data under custom metadata keys; '
        'in native C and C++ consumers, keys or values whose lengths exceed what the '
        'consuming code anticipates can trigger heap-based buffer overflows. The '
        'version field controls how the remainder of the file is parsed; supplying an '
        'unsupported version value can trigger undefined behavior in parsers without '
        'explicit version gating.')

    add_heading(doc, '3.5  Practical Attack Path', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The end-to-end attack path requires no novel infrastructure. The actor '
        'creates a Hugging Face account whose name differs from a popular publisher '
        'by a single character transposition. They publish a modified checkpoint '
        'whose malicious payload is embedded in a tensor never accessed during '
        'standard evaluation. Engineers performing searches encounter the forked '
        'repository, pull the artifact, and it clears all code-focused CI checks '
        'because no ML model inspection is configured. The artifact is registered in '
        'the internal model registry, promoted to staging, and eventually deployed '
        'to a production inference endpoint.')

    # ---- 4. System Architecture ----
    add_heading(doc, '4  SYSTEM ARCHITECTURE', font_size=11, space_before=12)

    add_heading(doc, '4.1  Design Principles', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The model-provenance-guard framework is grounded in four design principles '
        'derived from established supply chain security practice and the zero-trust '
        'architecture posture described by Mudusu and Gentyala [1].')
    for label, desc in [
        ('Shift-left enforcement.',
         ' Provenance checks must occur at the model intake boundary, before any '
         'artifact is registered in an internal model registry or promoted beyond '
         'the staging environment.'),
        ('Defense in depth.',
         ' No single control is sufficient. Hash verification, signature verification, '
         'and binary inspection each address a distinct portion of the threat surface.'),
        ('Hard failures, not warnings.',
         ' Any control failure must block pipeline progression. A gate that emits '
         'warnings while allowing promotion provides no practical security benefit.'),
        ('Registry integrity bootstrapping.',
         ' The trusted hash registry must itself be protected against tampering. '
         'The pipeline validates the registry\'s own SHA-256 hash before accepting '
         'any registry lookup result.'),
    ]:
        add_mixed(doc, [(label, True, False), (desc, False, False)], space_before=3)

    add_heading(doc, '4.2  Pipeline Architecture', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The provenance gate is implemented as a GitHub Actions workflow with five '
        'sequentially dependent jobs. Table 1 summarizes each job and its security '
        'objective.')

    # Pipeline table
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Job'
    hdr[1].text = 'Security Objective'
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    shade_table_header(table.rows[0])

    rows_data = [
        ('1. Registry Integrity Check',
         'Computes SHA-256 of trusted_models.json and compares against a sealed '
         'repository secret.'),
        ('2. Hash Verification',
         'Downloads artifact, computes SHA-256, and checks against the trusted '
         'registry, not the model card.'),
        ('3. Cosign Signature Verification',
         'Verifies .bundle or .sig file against an organization public key stored '
         'as a repository secret.'),
        ('4. Binary Format Inspection',
         'Routes to picklescan (PyTorch) or verify_weights.py (Safetensors/GGUF) '
         'based on file extension.'),
        ('5. Provenance Gate',
         'Evaluates all upstream results. Fails the pipeline if any required job '
         'did not succeed.'),
    ]
    for i, (job, obj) in enumerate(rows_data):
        row = table.rows[i + 1]
        row.cells[0].text = job
        row.cells[1].text = obj
    style_table(table)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run('Table 1: Five-stage provenance gate pipeline jobs.')
    r.italic = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'

    add_heading(doc, '4.3  Trusted Registry Design', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The trusted model registry is a JSON file maintained in the same repository '
        'as the workflow, protected by branch protection rules requiring at least one '
        'approved reviewer on any pull request that modifies it. Each entry records '
        'the model name, artifact filename, SHA-256 hash computed by a human reviewer '
        'at initial review time, source URL and commit hash, reviewing engineer\'s '
        'identity, review timestamp, and an optional Cosign bundle path.')
    add_body(doc,
        'A critical design decision is that the trusted hash must be computed by a '
        'human reviewer from the artifact at initial review time and must not be '
        'sourced from the model card or model hub metadata. An adversary with control '
        'over a model repository can update the published hash at any time. Only '
        'hashes computed independently and stored in a separately controlled registry '
        'can provide supply chain integrity guarantees.')

    # ---- 5. Implementation ----
    add_heading(doc, '5  IMPLEMENTATION', font_size=11, space_before=12)

    add_heading(doc, '5.1  Safetensors Header Inspector', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The inspect_safetensors() function performs seven checks before returning a '
        'pass or fail verdict, and in no case does it load tensor data into memory.')
    checks = [
        ('Check 1: Minimum file size.', ' A file smaller than eight bytes cannot be a valid Safetensors file and is rejected immediately.'),
        ('Check 2: Header size bounds.', ' The declared header size is validated against a configurable ceiling (100 MB by default) and a warning threshold (10 MB).'),
        ('Check 3: UTF-8 validity.', ' Header bytes must decode as valid UTF-8 before JSON parsing is attempted.'),
        ('Check 4: JSON validity.', ' The decoded header string must parse as valid JSON.'),
        ('Checks 5-6: Tensor entry validation.', ' Each tensor entry must contain a valid dtype, a well-formed shape field, and valid data_offsets.'),
        ('Check 7: Data offset bounds.', ' The maximum data offset declared across all tensor entries must not exceed the actual file size.'),
    ]
    for label, desc in checks:
        add_mixed(doc, [(label, True, False), (desc, False, False)], space_before=3)

    add_heading(doc, '5.2  GGUF Inspector', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The inspect_gguf() function validates the binary header and audits metadata '
        'key names without parsing metadata values. The function first checks that '
        'the file begins with the four magic bytes GGUF (ASCII). It then reads the '
        'version field, flagging versions outside the set {2, 3} as anomalous. It '
        'reads the declared tensor count and key-value count, flagging the latter if '
        'it exceeds 10,000. For each metadata key, the function reads the key length '
        'as a little-endian 64-bit integer, rejects keys longer than 1,024 bytes, '
        'validates UTF-8 encoding, and checks whether the key name begins with one '
        'of 28 known architecture-specific prefixes defined in the GGUF v3 '
        'specification.')

    add_heading(doc, '5.3  Cross-Platform Correctness', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The initial implementation hardcoded /tmp as the report output directory '
        'and used datetime.utcnow(), which was deprecated in Python 3.12. We '
        'corrected the report path to use tempfile.gettempdir(), which resolves '
        'correctly on Windows, Linux, and macOS, and replaced the deprecated call '
        'with datetime.now(datetime.timezone.utc).')

    add_heading(doc, '5.4  Hash Verification and Registry Lookup', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The verify_against_registry() function computes the SHA-256 digest of the '
        'artifact in 1 MB chunks to accommodate files that routinely exceed 4 GB, '
        'then queries the trusted registry for a matching entry by filename. If no '
        'entry is found, the function returns a permissive result with a warning, '
        'leaving the hard failure to the GitHub Actions pipeline step.')

    # ---- 6. Evaluation ----
    add_heading(doc, '6  EVALUATION', font_size=11, space_before=12)

    add_heading(doc, '6.1  Test Methodology', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'We constructed eight synthetic test cases covering the Safetensors and GGUF '
        'formats. Each test case programmatically constructs a binary artifact that '
        'embodies a specific structural condition and invokes the inspector with the '
        '--skip-hash flag to isolate binary inspection controls from the registry '
        'hash check. The expected exit code (0 for pass, 1 for failure) is compared '
        'against the actual exit code to determine test outcome.')

    add_heading(doc, '6.2  Results', font_size=10, bold=True, space_before=6)
    add_body(doc, 'Table 2 summarizes the outcomes across all eight test cases.')

    # Results table
    table2 = doc.add_table(rows=9, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Test Case'
    hdr2[1].text = 'Expected'
    hdr2[2].text = 'Result'
    for cell in hdr2:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    shade_table_header(table2.rows[0])

    results_data = [
        ('Safetensors: valid file (F32, correct offsets)', 'PASS', 'PASS'),
        ('Safetensors: invalid dtype (EVIL_TYPE)', 'FAIL', 'FAIL'),
        ('Safetensors: data offset past EOF', 'FAIL', 'FAIL'),
        ('Safetensors: non-JSON header bytes', 'FAIL', 'FAIL'),
        ('GGUF: valid file, zero KV entries', 'PASS', 'PASS'),
        ('GGUF: valid general.architecture key', 'PASS', 'PASS'),
        ('GGUF: wrong magic bytes (EVIL)', 'FAIL', 'FAIL'),
        ('GGUF: unsupported version (v99), non-strict mode', 'PASS', 'PASS'),
    ]
    for i, (case, exp, res) in enumerate(results_data):
        row = table2.rows[i + 1]
        row.cells[0].text = case
        row.cells[1].text = exp
        row.cells[2].text = res
    style_table(table2)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run('Table 2: Binary inspector evaluation results.')
    r.italic = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'

    add_body(doc,
        'All eight test cases produced the expected outcome, yielding a false '
        'negative rate of 0.0% and a false positive rate of 0.0% under the tested '
        'conditions. The unsupported GGUF version case correctly produces a warning '
        'annotation in the inspection report while allowing the artifact through, '
        'consistent with the design intent.',
        space_before=4)

    add_heading(doc, '6.3  Performance Characteristics', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'Binary inspection latency is bounded by file read bandwidth rather than '
        'computation. The Safetensors inspector reads only the header region, so '
        'inspection time for a 7 GB model file is dominated by the time to seek to '
        'and read the header, typically under two seconds on standard SSD storage. '
        'The GGUF inspector reads the fixed-length header and the first metadata key, '
        'a region always well under 1 KB, making its latency effectively constant '
        'regardless of file size. Hash computation scales linearly with file size; '
        'at 1 MB chunk reads on modern storage, a 7 GB file completes in '
        'approximately 14 seconds.')

    # ---- 7. Discussion ----
    add_heading(doc, '7  DISCUSSION', font_size=11, space_before=12)

    add_heading(doc, '7.1  Limitations', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The most consequential limitation of the framework is its inability to '
        'detect neural backdoors embedded in weight values. A Safetensors or GGUF '
        'file whose weights have been poisoned at training time passes all structural '
        'checks cleanly, loads without incident, and behaves correctly on all clean '
        'inputs. The malicious behavior surfaces only on inputs containing the '
        'attacker\'s trigger pattern [3]. Addressing this requires behavioral '
        'monitoring at inference time: continuous analysis of output distributions, '
        'activation statistics, and response latency. This is a complementary '
        'defense operating at the application layer and falls outside the scope '
        'of the present framework.')
    add_body(doc,
        'The GGUF inspector\'s metadata audit is currently limited to key name '
        'classification. Full value inspection requires a complete GGUF value-skip '
        'implementation. We recommend that teams requiring full metadata value '
        'inspection use gguf-py or llama.cpp\'s gguf-dump utility as a complement.',
        space_before=4)

    add_heading(doc, '7.2  Scope and Integration', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'The framework is designed to sit at the model intake boundary: between the '
        'artifact download step and the step that registers the artifact in an '
        'internal model registry. Organizations should treat this framework as one '
        'layer in a defense-in-depth strategy aligned with the NIST AI RMF [13] '
        'and the OWASP LLM Top 10 [16].')

    add_heading(doc, '7.3  Future Work', font_size=10, bold=True, space_before=6)
    add_body(doc,
        'Four extensions represent near-term priorities: (1) extending the binary '
        'inspector to cover ONNX protobuf models and TensorFlow SavedModel artifacts; '
        '(2) integrating AI Software Bill of Materials (AI-SBOM) generation into the '
        'provenance gate output [12, 13]; (3) implementing complete GGUF metadata '
        'value parsing to enable schema-level validation; (4) investigating '
        'integration with inference-time behavioral monitoring to close the neural '
        'backdoor detection gap.')

    # ---- 8. Conclusion ----
    add_heading(doc, '8  CONCLUSION', font_size=11, space_before=12)
    add_body(doc,
        'The machine learning supply chain is not merely a software supply chain. '
        'It is a binary supply chain in which the most sensitive components---weight '
        'files that encode model behavior---are distributed through infrastructure '
        'that provides no cryptographic integrity guarantees by default and consumed '
        'by tooling that offers no structural inspection prior to loading. This gap '
        'is neither obscure nor theoretical; it has been exploited in documented '
        'incidents and characterized in peer-reviewed vulnerability research across '
        'all three dominant weight file formats.')
    add_body(doc,
        'We have presented model-provenance-guard, a framework that addresses this '
        'gap with a combination of binary-level inspection and cryptographic '
        'provenance enforcement, deployed at the CI/CD boundary where prevention '
        'remains possible. Against eight synthetic adversarial test cases spanning '
        'the Safetensors and GGUF attack surfaces, the inspector returned zero false '
        'negatives and zero false positives.',
        space_before=4)
    add_body(doc,
        'The underlying principle is not novel: build the gate, not more alerts. '
        'Monitoring a threat class that cannot be observed through existing '
        'instrumentation generates no signal. The response to weaponized weights is '
        'structural, and it must be implemented before the next artifact a team '
        'stages for production introduces a threat that no existing tool was designed '
        'to see.',
        space_before=4)

    # ---- Acknowledgments ----
    add_heading(doc, 'ACKNOWLEDGMENTS', font_size=10, space_before=12)
    add_body(doc,
        'The authors thank the security research community for disclosing the '
        'llama.cpp heap-buffer-overflow vulnerabilities (Cisco Talos, 2024) and the '
        'Trail of Bits team for the Safetensors library security assessment (2023), '
        'both of which directly informed the threat model developed in this work. '
        'The model-provenance-guard toolkit is released under the MIT License and '
        'is available at https://github.com/sunilgentyala/model-provenance-guard.')

    # ---- References ----
    add_heading(doc, 'REFERENCES', font_size=10, space_before=12)
    refs = [
        '[1] Mudusu, S. K. and Gentyala, S. 2026. Zero-Trust Data Pipelines for AI Systems. Journal of Recent Trends in Computer Science and Engineering, 14(2), 10-25. DOI: 10.70589/JRTCSE.2026.14.2.2',
        '[2] Koishybayev, I. et al. 2025. PickleBall: Secure Deserialization of Pickle-based Machine Learning Models. In Proc. ACM CCS \'25. DOI: 10.1145/3719027.3765037',
        '[3] Gu, T., Dolan-Gavitt, B., and Garg, S. 2019. BadNets: Evaluating Backdooring Attacks on Deep Neural Networks. IEEE Access, 7, 47230-47244. DOI: 10.1109/ACCESS.2019.2909068',
        '[4] Jiang, J. et al. 2024. Models Are Codes: Towards Measuring Malicious Code Poisoning Attacks on Pre-trained Model Hubs. In Proc. ASE \'24. DOI: 10.1145/3691620.3695271',
        '[5] Communications of the ACM Staff. 2025. Malicious AI Models Undermine Software Supply-Chain Security. Commun. ACM. DOI: 10.1145/3704724',
        '[6] Newman, Z., Meyers, J. S., and Torres-Arias, S. 2022. Sigstore: Software Signing for Everybody. In Proc. ACM CCS \'22, 2353-2367. DOI: 10.1145/3548606.3560596',
        '[7] Schlegel, M. and Sattler, K.-U. 2024. Capturing End-to-End Provenance for Machine Learning Pipelines. Information Systems, 132, 102495. DOI: 10.1016/j.is.2024.102495',
        '[8] Dey, S. et al. 2025. An Empirical Study of Safetensors\' Usage Trends and Developers\' Perceptions. arXiv:2501.02170.',
        '[9] Jiang, W. et al. 2024. CryptoTensors: A Light-Weight LLM File Format for Highly-Secure Model Distribution. arXiv:2512.04580.',
        '[10] Anonymous. 2024. Mind the Gap: A Practical Attack on GGUF Quantization. OpenReview.',
        '[11] MITRE Corporation. 2023. MITRE ATLAS: Adversarial Threat Landscape for Artificial-Intelligence Systems. https://atlas.mitre.org',
        '[12] Gentyala, S. 2026. Securing the AI Supply Chain: A Framework for AI SBOMs. Transactions on ML and AI, 14(1), 119-129. DOI: 10.14738/tmlai.1401.19884',
        '[13] National Institute of Standards and Technology. 2023. AI Risk Management Framework (AI RMF 1.0). NIST AI 100-1. DOI: 10.6028/NIST.AI.100-1',
        '[14] MITRE Corporation. 2023. MITRE ATLAS. https://atlas.mitre.org',
        '[15] OpenSSF AI/ML Working Group. 2025. Launch of Model Signing v1.0. Open Source Security Foundation Blog.',
        '[16] OWASP Gen AI Security Project. 2025. LLM03:2025 Supply Chain. OWASP GenAI Security Top 10.',
        '[17] Gentyala, S. 2026. Operationalising AI Bills of Materials for Verifiable AI Provenance. Frontiers in Computer Science. DOI: 10.3389/fcomp.2026.1735919',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(ref)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9)

    out = os.path.join(OUTPUT_DIR, 'acm_paper.docx')
    doc.save(out)
    print(f'Saved: {out}')


# ---------------------------------------------------------------------------
# Elsevier Paper
# ---------------------------------------------------------------------------

def build_elsevier():
    doc = Document()

    # A4 page size
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    set_margins(doc, top=1.0, bottom=1.0, left=0.75, right=0.75)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    BLUE = RGBColor(0, 70, 127)
    FONT = 'Arial'

    # ---- Journal header ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Computers & Security')
    r.italic = True; r.font.name = FONT; r.font.size = Pt(10)
    r.font.color.rgb = BLUE

    add_rule(doc)

    # ---- Title ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(
        'Weaponized Weights: A Cryptographic Provenance and Binary Inspection '
        'Framework for Securing Machine Learning Model Artifacts in Enterprise '
        'MLOps Pipelines'
    )
    r.bold = True; r.font.name = FONT; r.font.size = Pt(17)
    r.font.color.rgb = BLUE

    # ---- Authors ----
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('Sunil Gentyala')
    r1.bold = True; r1.font.name = FONT; r1.font.size = Pt(12)
    r_sup = p.add_run('a,*')
    r_sup.font.name = FONT; r_sup.font.size = Pt(8)
    r_sup.font.superscript = True
    r2 = p.add_run(', Rakesh Prakash')
    r2.bold = True; r2.font.name = FONT; r2.font.size = Pt(12)
    r_sup2 = p.add_run('b')
    r_sup2.font.name = FONT; r_sup2.font.size = Pt(8)
    r_sup2.font.superscript = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('a HCLTech, Cybersecurity and AI Security Practice, Dallas, TX, USA')
    r.font.name = FONT; r.font.size = Pt(9)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('b University of Colorado Boulder, Department of Computer Science, Boulder, CO, USA')
    r.font.name = FONT; r.font.size = Pt(9)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('* Corresponding author. sunil.gentyala@ieee.org | Rakesh.Prakash@colorado.edu')
    r.font.name = FONT; r.font.size = Pt(9); r.italic = True

    add_rule(doc)

    # ---- Article info + keywords ----
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run('Keywords: ')
    r.bold = True; r.font.name = FONT; r.font.size = Pt(9)
    r2 = p.add_run(
        'Machine learning security; Supply chain attack; Model provenance; '
        'Safetensors; GGUF; Pickle deserialization; MLOps security; '
        'Cryptographic verification; CI/CD security; Binary inspection'
    )
    r2.font.name = FONT; r2.font.size = Pt(9)

    add_rule(doc)
    set_two_columns(doc)

    # ---- Abstract ----
    add_heading(doc, 'Abstract', font_name=FONT, font_size=11, space_before=8)
    add_body(doc,
        'The widespread adoption of publicly distributed machine learning model '
        'artifacts has introduced a structural security risk that conventional '
        'enterprise tooling is fundamentally unable to address. Weight files '
        'serialized in PyTorch checkpoint, Safetensors, and GGUF formats are '
        'routinely downloaded from public model repositories and admitted into '
        'production inference pipelines without binary-level structural inspection '
        'or cryptographic provenance verification. Existing static analysis, dynamic '
        'analysis, and antivirus infrastructure have no visibility into these binary '
        'formats, leaving organizations exposed to a class of supply chain attack '
        'that operates entirely below the application layer.',
        font_name=FONT, space_before=4)
    add_body(doc,
        'This paper presents model-provenance-guard, a production-ready, open-source '
        'toolkit comprising a Python binary inspector (verify_weights.py) and a '
        'five-stage GitHub Actions provenance gate (model_scan.yml). The pipeline '
        'enforces SHA-256 hash verification against an internal trusted registry, '
        'Sigstore/Cosign cryptographic signature checking, pickle opcode scanning, '
        'Safetensors header anomaly detection, and GGUF metadata key auditing before '
        'any artifact is promoted within a downstream pipeline. We evaluate the '
        'toolkit against eight synthetic adversarial test cases covering dtype '
        'manipulation, data offset injection, magic byte substitution, malformed '
        'JSON header attacks, and unsupported version fields. The binary inspector '
        'achieved a false negative rate of 0.0% and a false positive rate of 0.0% '
        'across all test cases, demonstrating that shift-left provenance enforcement '
        'is both technically sound and operationally practical at the CI/CD boundary. '
        'The framework is released as open-source software at '
        'https://github.com/sunilgentyala/model-provenance-guard.',
        font_name=FONT, space_before=4)

    def H1(text): add_heading(doc, text, font_name=FONT, font_size=12, space_before=12, color=BLUE)
    def H2(text): add_heading(doc, text, font_name=FONT, font_size=11, bold=True, space_before=8)
    def Body(text, **kw): add_body(doc, text, font_name=FONT, **kw)

    # ---- 1. Introduction ----
    H1('1. Introduction')
    Body(
        'The security research community\'s focus on artificial intelligence systems '
        'has been concentrated primarily on prompt injection, model output '
        'manipulation, and inference-time adversarial inputs. Those threats are real '
        'and warrant continued attention. However, their prominence in the discourse '
        'has permitted a structurally more fundamental threat to develop with '
        'insufficient scrutiny. Adversaries with the means and motivation to compromise '
        'enterprise machine learning deployments are no longer limited to crafting '
        'clever inputs at runtime. They are embedding threats directly inside compiled '
        'weight artifacts that machine learning engineers download from public model '
        'hubs as a routine operational act, with no inspection beyond confirming that '
        'the download completed successfully.')
    Body(
        'The attack surface has shifted from the application layer to the binary '
        'layer. Static application security testing tools analyze source code. '
        'Dynamic testing tools probe running applications. Antivirus engines match '
        'known malware signatures against file content. None of these capabilities '
        'has any meaningful visibility into the internal structure of a 4-bit '
        'quantized GGUF binary, a Safetensors archive, or a PyTorch checkpoint '
        'serialized with Python\'s pickle protocol. When an engineer issues a '
        'huggingface-cli download command and stages the resulting artifact in a '
        'corporate MLOps pipeline, no system in the standard enterprise security '
        'stack has examined what is structurally present inside that binary.',
        space_before=4)
    Body(
        'Three weight file formats account for the substantial majority of model '
        'artifacts in contemporary enterprise deployment. PyTorch .pt and .pth '
        'checkpoints rely on Python\'s pickle serialization protocol, which by '
        'design permits arbitrary Python objects to be deserialized, creating a '
        'direct remote code execution pathway at the moment torch.load() is called '
        '[2]. Safetensors was designed specifically to eliminate deserialization-time '
        'code execution; however, its JSON header parsing stage introduces attack '
        'surfaces through parser exploitation, metadata field abuse, and the silent '
        'embedding of backdoored weights [8]. GGUF, the dominant format for '
        'quantized local inference, includes a flexible metadata section with no '
        'schema enforcement, and multiple heap-buffer-overflow vulnerabilities were '
        'disclosed in its reference C parser in 2024 [11].',
        space_before=4)
    Body(
        'This paper makes the following contributions: '
        '(1) A structured, format-level threat model covering the attack surfaces of '
        'PyTorch checkpoints, Safetensors archives, and GGUF binaries. '
        '(2) verify_weights.py, a Python binary inspector performing Safetensors '
        'header anomaly detection and GGUF magic-byte and metadata key auditing '
        'without loading tensor data into memory. '
        '(3) A five-stage GitHub Actions provenance gate integrating registry hash '
        'verification, Sigstore/Cosign signature verification, picklescan opcode '
        'analysis, and binary inspection into a blocking CI/CD workflow. '
        '(4) An empirical evaluation against eight synthetic adversarial test cases '
        'demonstrating zero false negatives and zero false positives across all '
        'tested attack scenarios. '
        '(5) Cross-platform portability corrections and deprecation fixes applied '
        'to the reference implementation.',
        space_before=4)

    # ---- 2. Background ----
    H1('2. Background and Related Work')

    H2('2.1. Machine Learning Supply Chain Attacks')
    Body(
        'Supply chain attacks against software artifacts have been documented for '
        'over a decade, but their application to machine learning weight files is '
        'a recent and rapidly evolving phenomenon. Jiang et al. [4] conducted a '
        'systematic measurement study of malicious code poisoning attacks on '
        'pre-trained model hubs, documenting how attackers exploit namespace trust '
        'and version ambiguity to deliver weaponized artifacts to downstream '
        'consumers. The Communications of the ACM survey [5] argues that the '
        'combination of billions of monthly model downloads and weak provenance '
        'controls makes model hubs uniquely attractive targets relative to '
        'traditional package registries.')
    Body(
        'The zero-trust model for AI data pipelines articulated by Mudusu and '
        'Gentyala [1] establishes that the never-trust-always-verify principle '
        'must extend directly to data and model artifacts, not merely to network '
        'perimeters. That work demonstrates that cryptographic authentication '
        'coupled with tamper-evident lineage tracking achieves complete detection '
        'of artifact tampering and component impersonation in production AI systems, '
        'validating the architectural approach adopted in the present framework.',
        space_before=4)

    H2('2.2. Pickle Deserialization Vulnerabilities')
    Body(
        'The most thoroughly characterized vulnerability class in the ML weight '
        'file domain is Python pickle deserialization. The pickle protocol\'s '
        '__reduce__ mechanism enables arbitrary Python callables to execute during '
        'deserialization, providing a direct code execution pathway that requires '
        'no user interaction beyond a standard torch.load() call [2]. Koishybayev '
        'et al. [2] present PickleBall, a secure deserialization system that '
        'intercepts pickle opcodes prior to object reconstruction and enforces a '
        'safe-by-default allowlist, confirming that the vulnerability is technically '
        'addressable but requires purpose-built defense rather than general-purpose '
        'antivirus.')

    H2('2.3. Safetensors Format Security')
    Body(
        'The Safetensors format [8] was developed specifically to eliminate '
        'deserialization-time code execution by replacing the pickle protocol with '
        'a JSON-encoded header followed by raw tensor byte data. The Trail of Bits '
        'security assessment of the reference implementation (2023) confirmed the '
        'absence of deserialization vulnerabilities. However, the JSON header '
        'parsing stage is not without risk: oversized headers, embedded null bytes, '
        'and deeply nested JSON structures can exploit implementation-specific '
        'parser behavior. CryptoTensors [9] proposes an extension that adds '
        'tensor-level encryption and embedded access control policies.')

    H2('2.4. GGUF Format and Binary Parsing Vulnerabilities')
    Body(
        'GGUF, introduced in August 2023 and adopted as the dominant distribution '
        'format for quantized inference in Ollama and llama.cpp, presents a binary '
        'attack surface that differs qualitatively from the other two formats. '
        'Cisco Talos disclosed three heap-buffer-overflow vulnerabilities in the '
        'llama.cpp GGUF parser in early 2024, all arising from insufficient bounds '
        'checking in the metadata key-value section parsing logic. A separate '
        'research contribution demonstrates that quantization error in GGUF models '
        'provides sufficient numerical flexibility to construct adversarial models '
        'whose behavior is benign under full-precision analysis but malicious under '
        'quantized inference [11].')

    H2('2.5. Neural Backdoor Attacks')
    Body(
        'The neural backdoor threat was formalized by Gu et al. in the BadNets '
        'paper [3], which demonstrated that a model can be trained to behave '
        'correctly on all clean inputs while triggering targeted misclassification '
        'on inputs containing a specific attacker-controlled pattern. This attack '
        'class is undetectable by any binary inspection technique: the weight file '
        'loads cleanly, passes all format-level checks, and produces correct outputs '
        'during standard evaluation. Detection requires behavioral monitoring at '
        'inference time, specifically analysis of output probability distributions '
        'and internal activation statistics against baselines established from '
        'known-clean model weights.')

    H2('2.6. Software Artifact Signing')
    Body(
        'Sigstore [6] provides a keyless signing infrastructure built on ephemeral '
        'OIDC identity tokens, substantially lowering the adoption barrier relative '
        'to traditional GPG-based signing schemes. The Cosign component of the '
        'Sigstore ecosystem supports signing and verification of arbitrary binary '
        'objects, making it directly applicable to model weight files. As of 2025, '
        'Sigstore underlies artifact signing for npm, PyPI, and Kubernetes. The '
        'OpenSSF AI/ML Working Group has published a dedicated model signing '
        'specification [15] that adapts the Sigstore framework specifically to the '
        'machine learning weight file distribution context.')

    H2('2.7. ML Artifact Provenance')
    Body(
        'Schlegel and Sattler [7] present a PROV-compliant provenance model for '
        'end-to-end ML pipelines that captures artifact lineage through MLflow and '
        'Git activities. Gentyala [12] frames the AI supply chain problem in terms '
        'of AI Software Bills of Materials (AI-SBOMs), demonstrating that CycloneDX '
        'and SPDX standards can be adapted to encode model weight metadata and '
        'cryptographic attestations within familiar SBOM infrastructure. The NIST '
        'AI Risk Management Framework [13] and MITRE ATLAS [14] provide '
        'complementary governance and threat taxonomy infrastructure, and the '
        'OWASP LLM Top 10 explicitly classifies supply chain compromise as a '
        'critical risk category for LLM deployments [16].')

    # ---- 3. Threat Model ----
    H1('3. Threat Model')

    H2('3.1. Adversary Model')
    Body(
        'We model an adversary who has obtained write access to a model repository '
        'on a public model hub. Three realistic entry points exist: account '
        'compromise of a legitimate model publisher, a pull request containing a '
        'malicious weight modification merged by an inattentive maintainer, and '
        'creation of a confusingly named fork designed to capture traffic from '
        'engineers who do not verify repository provenance before downloading. '
        'The adversary\'s objective is to deliver a weaponized weight artifact to '
        'at least one GPU-enabled inference endpoint inside a target enterprise '
        'environment. The adversary has read access to the model hub API and can '
        'observe download statistics and model version popularity. We do not assume '
        'the adversary has prior network access to the target organization\'s '
        'internal infrastructure.')

    H2('3.2. Attack Surface: PyTorch Checkpoints')
    Body(
        'PyTorch .pt and .pth files use Python\'s pickle serialization protocol. '
        'The pickle module supports the __reduce__ protocol, which specifies a '
        'callable and its arguments to be used during deserialization. An adversary '
        'can embed an object whose __reduce__ method encodes a system command, an '
        'outbound network connection, or a persistence mechanism. This payload '
        'executes synchronously during the call to torch.load(), before the caller '
        'can inspect the returned object. No user action beyond the standard model '
        'loading call is required. On inference infrastructure with GPU access and '
        'network egress, this constitutes a high-value beachhead.')

    H2('3.3. Attack Surface: Safetensors')
    Body(
        'Safetensors eliminates deserialization-time code execution but introduces '
        'three distinct residual attack surfaces. First, the JSON header that is '
        'parsed before any tensor data is accessed is vulnerable to parser '
        'exploitation through oversized payloads, malformed Unicode sequences, and '
        'JSON structures designed to trigger recursion-depth failures or memory '
        'exhaustion in consuming applications. Second, the __metadata__ dictionary '
        'accepts arbitrary string key-value pairs without schema enforcement; '
        'applications that log or forward metadata values without sanitization may '
        'expose themselves to secondary injection attacks. Third, Safetensors files '
        'can carry weights modified through backdoor training attacks [3]; these '
        'files pass all structural checks and behave correctly on clean inputs.')

    H2('3.4. Attack Surface: GGUF')
    Body(
        'The GGUF metadata section has no schema enforcement in the format '
        'specification. An adversary can embed arbitrary data under custom metadata '
        'keys; in native C and C++ consumers, keys or values that exceed what the '
        'consuming code anticipates can trigger heap-based buffer overflows. The '
        'version field at bytes 4-7 controls how the remainder of the file is '
        'parsed; supplying a non-standard version value can induce undefined parsing '
        'behavior in consumers that do not implement explicit version gating. The '
        'tensor count field at bytes 8-15 is trusted by consuming libraries; a '
        'mismatch between the declared count and the number of tensors actually '
        'present in the file causes the parser to over-read its buffer.')

    H2('3.5. Practical Attack Path')
    Body(
        'A threat actor creates a Hugging Face account whose name differs from a '
        'popular model publisher\'s account by a single character. They publish a '
        'modified checkpoint with a plausible model card, including a hash value '
        'they compute themselves after modification. Engineers search for the target '
        'model, encounter the forked repository, and download the artifact. The '
        'artifact clears all code-focused CI checks because no binary model '
        'inspection exists in the pipeline. It is registered in the internal model '
        'registry, promoted to staging, and eventually deployed to a production '
        'inference endpoint.')

    # ---- 4. System Architecture ----
    H1('4. System Architecture')

    H2('4.1. Design Principles')
    Body(
        'The model-provenance-guard framework is built on four design principles '
        'grounded in supply chain security practice and the zero-trust posture '
        'for AI systems described by Mudusu and Gentyala [1].')
    for label, desc in [
        ('Shift-left enforcement.', ' Security checks must occur at the model intake boundary, before any artifact is registered in an internal model registry or promoted beyond staging. Post-deployment inspection provides no prevention value.'),
        ('Defense in depth.', ' No single control addresses the full threat surface. Hash verification, signature verification, and binary inspection each address a distinct portion of the threat surface independently.'),
        ('Hard failure semantics.', ' Any control failure must block pipeline progression. A gate that emits warnings while continuing to promote artifacts provides compliance theater without security value.'),
        ('Registry integrity bootstrapping.', ' The trusted hash registry must itself be protected against tampering. The pipeline verifies the registry\'s own SHA-256 hash against a sealed repository secret as its first action.'),
    ]:
        add_mixed(doc, [(label, True, False), (desc, False, False)],
                  font_name=FONT, space_before=3)

    H2('4.2. Five-Stage Provenance Gate')
    Body('Table 1 describes the five GitHub Actions jobs that constitute the provenance gate.')

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Job'
    hdr[1].text = 'Security Objective'
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    shade_table_header(table.rows[0])
    rows_data = [
        ('1. Registry Integrity',
         'Verifies SHA-256 of trusted_models.json against a sealed repository secret before any registry lookup is trusted.'),
        ('2. Hash Verification',
         'Computes artifact SHA-256 immediately after download and compares against the internal trusted registry, not the model card.'),
        ('3. Cosign Signature',
         'Verifies a Cosign .bundle or .sig file against an organization public key stored as a sealed repository secret.'),
        ('4. Binary Inspection',
         'Routes to picklescan for PyTorch checkpoints, or to verify_weights.py for Safetensors and GGUF artifacts.'),
        ('5. Provenance Gate',
         'Evaluates all upstream job results and blocks the pipeline if any required job did not succeed. Writes a summary for audit logging.'),
    ]
    for i, (job, obj) in enumerate(rows_data):
        row = table.rows[i + 1]
        row.cells[0].text = job
        row.cells[1].text = obj
    style_table(table, font_name=FONT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run('Table 1: GitHub Actions provenance gate jobs and their security objectives.')
    r.italic = True; r.font.size = Pt(9); r.font.name = FONT

    H2('4.3. Trusted Registry Design')
    Body(
        'The trusted model registry is a version-controlled JSON file whose schema '
        'requires, for each registered artifact, the model name, the artifact '
        'filename, a SHA-256 hash computed by a human reviewer from the artifact '
        'at initial review time, the source URL and source repository commit hash, '
        'the reviewer\'s identity, the review timestamp, and an optional Cosign '
        'bundle path. The registry is protected by branch protection rules requiring '
        'at least one approved reviewer on any pull request that modifies it.')
    Body(
        'The critical design constraint is that the trusted hash must be computed '
        'independently by the reviewing engineer from the artifact directly, not '
        'sourced from the model card or model hub metadata API. An adversary with '
        'control over a model repository can update the model card hash at any time '
        'without triggering a new download. Only hashes computed independently and '
        'stored in a separately controlled registry can provide meaningful supply '
        'chain integrity guarantees.',
        space_before=4)

    # ---- 5. Implementation ----
    H1('5. Implementation')

    H2('5.1. Safetensors Header Inspector')
    Body(
        'The inspect_safetensors() function performs seven sequential checks before '
        'returning a pass or fail verdict. In no case does it load tensor data into '
        'memory; all checks operate exclusively on the header region.')
    checks = [
        ('Check 1: Minimum file size.', ' A file smaller than eight bytes cannot be a valid Safetensors file and is rejected immediately.'),
        ('Check 2: Header size bounds.', ' The declared header size is validated against a hard ceiling of 100 MB and a soft warning threshold of 10 MB.'),
        ('Check 3: UTF-8 validity.', ' Header bytes are decoded as UTF-8 before any JSON parsing is attempted, preventing a class of attacks exploiting malformed Unicode input.'),
        ('Check 4: JSON validity.', ' The decoded header must parse as valid JSON. A malformed JSON header constitutes a hard failure.'),
        ('Checks 5-6: Tensor entry field validation.', ' Each tensor entry must contain a valid dtype from the ten permitted Safetensors data types, a well-formed non-negative integer list shape field, and valid data_offsets.'),
        ('Check 7: Data offset bounds.', ' The maximum declared data offset must not exceed the actual file size, catching truncated files and phantom data region attacks.'),
    ]
    for label, desc in checks:
        add_mixed(doc, [(label, True, False), (desc, False, False)],
                  font_name=FONT, space_before=3)

    H2('5.2. GGUF Binary Inspector')
    Body(
        'The inspect_gguf() function validates the binary header and audits metadata '
        'key names. It does not parse metadata values, which would require a complete '
        'GGUF parser and would introduce its own attack surface. The function '
        'validates that the file begins with the four-byte GGUF magic sequence. '
        'It reads the version field (uint32, little-endian) and flags versions outside '
        'the supported set {2, 3} as anomalous. It reads the declared tensor count '
        'and key-value count fields, flagging counts above 10,000 as anomalous. For '
        'each declared metadata key, it reads the key length as a uint64, rejects '
        'keys longer than 1,024 bytes, validates UTF-8 encoding, and checks whether '
        'the key name begins with one of 28 known architecture-specific prefixes '
        'defined in the GGUF v3 specification. The presence of GGUF_STRICT=1 in the '
        'environment promotes unexpected key classifications from warnings to hard '
        'failures.')

    H2('5.3. Cross-Platform Corrections')
    Body(
        'The initial implementation contained two portability defects identified '
        'during evaluation. First, the write_report() function hardcoded /tmp as the '
        'report output directory, which is unavailable on Windows. This was corrected '
        'to use tempfile.gettempdir(), which resolves to the appropriate system '
        'temporary directory on Windows, Linux, and macOS. Second, both timestamp '
        'fields used datetime.utcnow(), which was deprecated in Python 3.12 and is '
        'scheduled for removal. These were replaced with '
        'datetime.now(datetime.timezone.utc), which returns a timezone-aware UTC '
        'datetime object compatible with Python 3.12 and later. Both defects would '
        'have caused the inspector to raise an unhandled exception after completing '
        'all inspection checks successfully, causing false failures on Windows '
        'environments.')

    H2('5.4. Hash Verification')
    Body(
        'The verify_against_registry() function computes SHA-256 in 1 MB chunks to '
        'accommodate model files that routinely exceed 4 GB. The digest is compared '
        'against the registry entry for the artifact\'s filename. If no entry is '
        'found, the function returns a permissive result and emits a warning; the '
        'GitHub Actions pipeline is configured to treat an unregistered artifact '
        'as a blocking failure at the workflow level, providing defense in depth '
        'against edge cases where the script-level permissive return would otherwise '
        'allow an unreviewed artifact through.')

    # ---- 6. Evaluation ----
    H1('6. Evaluation')

    H2('6.1. Test Methodology')
    Body(
        'We constructed eight synthetic binary artifacts, each representing either '
        'a structurally valid example or a specific attack scenario, and invoked '
        'the binary inspector against each using the --skip-hash flag to isolate '
        'the format inspection controls from the registry hash check. The expected '
        'exit code (0 for a passing verdict, 1 for a failing verdict) was compared '
        'against the actual exit code to determine the test outcome. The test cases '
        'were generated programmatically using Python\'s struct and json modules to '
        'construct precise binary artifacts with controlled structural properties.')

    H2('6.2. Test Cases and Results')
    Body('Table 2 presents the eight test cases, the expected outcome for each, and the outcome produced by the inspector. All eight cases produced the expected outcome.')

    table2 = doc.add_table(rows=9, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Test Case Description'
    hdr2[1].text = 'Expected'
    hdr2[2].text = 'Result'
    for cell in hdr2:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    shade_table_header(table2.rows[0])
    results_data = [
        ('Safetensors: valid F32 tensor, correct offsets', 'PASS', 'PASS'),
        ('Safetensors: invalid dtype (EVIL_TYPE)', 'FAIL', 'FAIL'),
        ('Safetensors: data offset beyond EOF', 'FAIL', 'FAIL'),
        ('Safetensors: non-JSON header bytes', 'FAIL', 'FAIL'),
        ('GGUF: valid file with zero metadata KV entries', 'PASS', 'PASS'),
        ('GGUF: valid general.architecture key', 'PASS', 'PASS'),
        ('GGUF: wrong magic bytes (EVIL)', 'FAIL', 'FAIL'),
        ('GGUF: version 99 (non-strict mode)', 'PASS', 'PASS'),
    ]
    for i, (case, exp, res) in enumerate(results_data):
        row = table2.rows[i + 1]
        row.cells[0].text = case
        row.cells[1].text = exp
        row.cells[2].text = res
    style_table(table2, font_name=FONT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run('Table 2: Binary inspector evaluation results across eight adversarial and benign test cases.')
    r.italic = True; r.font.size = Pt(9); r.font.name = FONT

    Body(
        'The false negative rate across all adversarial cases was 0.0%. The false '
        'positive rate across all benign cases was 0.0%. The unsupported GGUF '
        'version test confirmed that anomalous but structurally sound files pass '
        'in non-strict mode while generating a warning annotation in the inspection '
        'report, consistent with the framework\'s design intent.',
        space_before=4)

    H2('6.3. Performance Analysis')
    Body(
        'Binary inspection latency is bounded by file I/O bandwidth rather than '
        'computation. The Safetensors inspector reads only the header region '
        '(typically well under 1 MB), making inspection time effectively independent '
        'of model file size beyond the initial file open operation. The GGUF '
        'inspector reads the 24-byte fixed header and the first metadata key, a '
        'total read of under 1 KB under all circumstances. SHA-256 computation '
        'scales linearly with file size; at 1 MB chunk reads on standard SSD '
        'storage, a 7 GB model file completes in approximately 14 seconds.')

    # ---- 7. Discussion ----
    H1('7. Discussion')

    H2('7.1. Neural Backdoor Detection Gap')
    Body(
        'The most significant limitation of the framework is its inability to detect '
        'neural backdoors embedded in tensor weights. A model whose parameters have '
        'been poisoned at training time passes all structural checks correctly, loads '
        'without incident in any consuming framework, and produces correct outputs '
        'on all clean test inputs. The malicious behavior surfaces only on inputs '
        'containing the attacker\'s specific trigger pattern [3]. Binary inspection '
        'cannot address this threat class because the malicious content is numerically '
        'indistinguishable from legitimate model weights.')
    Body(
        'Addressing this gap requires behavioral monitoring at inference time: '
        'continuous analysis of output probability distributions, internal activation '
        'statistics, and response latency against baselines established from '
        'known-clean model weights. This defense is complementary to binary '
        'inspection and must be implemented as a separate layer in the '
        'defense-in-depth stack aligned with the NIST AI RMF [13] and '
        'MITRE ATLAS [14].',
        space_before=4)

    H2('7.2. GGUF Partial Metadata Coverage')
    Body(
        'The GGUF inspector\'s metadata audit is currently limited to key name '
        'classification. Full value inspection---validating that metadata values '
        'have the types, lengths, and content consistent with the declared model '
        'architecture---requires a complete GGUF value-skip implementation that '
        'correctly advances the file pointer across all GGUF value types. This is '
        'a non-trivial implementation task, since incorrect value-skip logic would '
        'itself constitute a parsing vulnerability. We recommend that teams requiring '
        'full metadata value inspection use gguf-py or llama.cpp\'s gguf-dump '
        'utility alongside verify_weights.py.')

    H2('7.3. Integration Context')
    Body(
        'The framework is designed as a gate at the model intake boundary between '
        'artifact download and internal model registry registration. It does not '
        'address security controls at earlier stages (dataset provenance, training '
        'environment integrity) or at later stages (inference-time behavioral '
        'monitoring, model output auditing). Organizations should position '
        'model-provenance-guard as one layer within a broader AI governance strategy '
        'aligned with the AI-SBOM framework described by Gentyala [12, 17].')

    H2('7.4. Future Work')
    Body(
        'Four directions represent the highest-priority extensions. First, extending '
        'verify_weights.py to support ONNX protobuf and TensorFlow SavedModel '
        'formats, which introduce distinct format-specific attack surfaces not covered '
        'by the current implementation. Second, integrating AI-SBOM generation into '
        'the provenance gate output, enabling downstream systems to consume '
        'cryptographically attested model lineage records in CycloneDX or SPDX '
        'format [12]. Third, implementing complete GGUF metadata value parsing to '
        'support schema-level validation. Fourth, investigating integration with '
        'inference-time behavioral monitoring to close the neural backdoor detection '
        'gap.')

    # ---- 8. Conclusion ----
    H1('8. Conclusion')
    Body(
        'Machine learning model deployment is, at its core, a binary artifact '
        'distribution problem. Weight files encoding model behavior are sourced from '
        'public infrastructure that provides no default cryptographic integrity '
        'guarantees and consumed by tooling that performs no structural inspection '
        'prior to loading. This gap is not a theoretical concern. It has been '
        'exploited in documented incidents, characterized in peer-reviewed '
        'vulnerability disclosures, and confirmed by multiple independent security '
        'research groups across all three dominant weight file formats.')
    Body(
        'This paper presents model-provenance-guard, a framework that closes this '
        'gap at the CI/CD boundary through the combination of binary-level inspection '
        'and cryptographic provenance enforcement. The five-stage GitHub Actions '
        'pipeline enforces defense in depth: registry integrity verification, '
        'SHA-256 hash validation, Cosign signature checking, opcode scanning, and '
        'binary format inspection operate as sequential, blocking gates before any '
        'artifact reaches an internal model registry. Against eight adversarial test '
        'cases spanning the Safetensors and GGUF attack surfaces, the inspector '
        'returned zero false negatives and zero false positives.',
        space_before=4)
    Body(
        'The operational recommendation is simple: build the gate before the next '
        'model artifact enters production. Monitor for a threat you can detect; '
        'inspect for a threat that must be caught before it loads.',
        space_before=4)

    # ---- Acknowledgments ----
    H1('Acknowledgments')
    Body(
        'The authors thank the Cisco Talos security research team for their '
        'disclosure of the llama.cpp GGUF parser heap-buffer-overflow vulnerabilities '
        '(2024) and the Trail of Bits team for the Safetensors library security '
        'assessment (2023), both of which directly informed the threat model '
        'developed in this work. The model-provenance-guard toolkit is released '
        'under the MIT License and is available at '
        'https://github.com/sunilgentyala/model-provenance-guard.')

    # ---- Competing interests ----
    H1('Declaration of Competing Interest')
    Body(
        'The authors declare that they have no known competing financial interests '
        'or personal relationships that could have appeared to influence the work '
        'reported in this paper.')

    # ---- Data availability ----
    H1('Data Availability')
    Body(
        'The model-provenance-guard toolkit, including all source code, test cases, '
        'and the GitHub Actions workflow, is publicly available at '
        'https://github.com/sunilgentyala/model-provenance-guard under the MIT '
        'License. No proprietary datasets were used in this study.')

    # ---- References ----
    H1('References')
    refs = [
        '[1] S. K. Mudusu and S. Gentyala, "Zero-Trust Data Pipelines for AI Systems: A Framework for Secure, Verifiable, and Auditable Data Engineering," J. Recent Trends Comput. Sci. Eng., vol. 14, no. 2, pp. 10-25, 2026, doi: 10.70589/JRTCSE.2026.14.2.2.',
        '[2] I. Koishybayev et al., "PickleBall: Secure Deserialization of Pickle-based Machine Learning Models," in Proc. ACM CCS \'25, 2025, doi: 10.1145/3719027.3765037.',
        '[3] T. Gu, B. Dolan-Gavitt, and S. Garg, "BadNets: Evaluating Backdooring Attacks on Deep Neural Networks," IEEE Access, vol. 7, pp. 47230-47244, 2019, doi: 10.1109/ACCESS.2019.2909068.',
        '[4] J. Jiang et al., "Models Are Codes: Towards Measuring Malicious Code Poisoning Attacks on Pre-trained Model Hubs," in Proc. ASE \'24, 2024, doi: 10.1145/3691620.3695271.',
        '[5] Commun. ACM Staff, "Malicious AI Models Undermine Software Supply-Chain Security," Commun. ACM, 2025, doi: 10.1145/3704724.',
        '[6] Z. Newman, J. S. Meyers, and S. Torres-Arias, "Sigstore: Software Signing for Everybody," in Proc. ACM CCS \'22, pp. 2353-2367, 2022, doi: 10.1145/3548606.3560596.',
        '[7] M. Schlegel and K.-U. Sattler, "Capturing End-to-End Provenance for Machine Learning Pipelines," Inf. Syst., vol. 132, p. 102495, 2024, doi: 10.1016/j.is.2024.102495.',
        '[8] S. Dey et al., "An Empirical Study of Safetensors\' Usage Trends and Developers\' Perceptions," arXiv:2501.02170, 2025.',
        '[9] W. Jiang et al., "CryptoTensors: A Light-Weight LLM File Format for Highly-Secure Model Distribution," arXiv:2512.04580, 2024.',
        '[10] Anonymous, "Mind the Gap: A Practical Attack on GGUF Quantization," OpenReview, 2024.',
        '[11] MITRE Corporation, "MITRE ATLAS: Adversarial Threat Landscape for Artificial-Intelligence Systems," 2023. [Online]. Available: https://atlas.mitre.org.',
        '[12] S. Gentyala, "Securing the AI Supply Chain: A Framework for AI Software Bills of Materials and Model Provenance Assurance," Trans. Mach. Learn. Artif. Intell., vol. 14, no. 1, pp. 119-129, 2026, doi: 10.14738/tmlai.1401.19884.',
        '[13] NIST, "Artificial Intelligence Risk Management Framework (AI RMF 1.0)," NIST AI 100-1, 2023, doi: 10.6028/NIST.AI.100-1.',
        '[14] MITRE Corporation, "MITRE ATLAS," 2023. [Online]. Available: https://atlas.mitre.org.',
        '[15] OpenSSF AI/ML Working Group, "Launch of Model Signing v1.0: Securing the Machine Learning Supply Chain," Open Source Security Foundation Blog, 2025.',
        '[16] OWASP Gen AI Security Project, "LLM03:2025 Supply Chain," OWASP GenAI Security Top 10, 2025.',
        '[17] S. Gentyala, "Operationalising Artificial Intelligence Bills of Materials for Verifiable AI Provenance and Lifecycle Assurance," Front. Comput. Sci., 2026, doi: 10.3389/fcomp.2026.1735919.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(ref)
        r.font.name = FONT
        r.font.size = Pt(9)

    out = os.path.join(OUTPUT_DIR, 'elsevier_paper.docx')
    doc.save(out)
    print(f'Saved: {out}')


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    build_acm()
    build_elsevier()
    print('Done.')
